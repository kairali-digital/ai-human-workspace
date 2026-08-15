#!/usr/bin/env python3
"""Build the printable beginner guide from its reviewed Markdown source."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "EVERYONE-ELSE-AI-HUMAN-HOMEWORK-GUIDE.md"
OUTPUT = ROOT / "EVERYONE-ELSE-AI-HUMAN-HOMEWORK-GUIDE.docx"

# compact_reference_guide preset, resolved numerically.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_DISTANCE = Inches(0.492)
FOOTER_DISTANCE = Inches(0.492)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GOLD = RGBColor(0x9A, 0x6B, 0x1F)
GRAY = RGBColor(0x55, 0x55, 0x55)
MUTED = RGBColor(0x6B, 0x72, 0x80)
INK = RGBColor(0x1F, 0x29, 0x37)
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
PALE_GRAY = "F4F6F9"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in margins.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_table_headers(doc):
    """Expose every designed table's first row to assistive technology."""
    for table in doc.tables:
        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        header = tr_pr.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            tr_pr.append(header)
        header.set(qn("w:val"), "true")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def next_numbering_id(numbering):
    ids = []
    for tag in ("w:abstractNum", "w:num"):
        attr = "w:abstractNumId" if tag.endswith("abstractNum") else "w:numId"
        for node in numbering.findall(qn(tag)):
            value = node.get(qn(attr))
            if value and value.isdigit():
                ids.append(int(value))
    return max(ids, default=0) + 1


def add_numbering_definition(doc, kind, start_value=1):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering)
    num_id = abstract_id + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_value))
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def style_paragraph(paragraph, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_together = False
    for run in paragraph.runs:
        set_font(run, size=11, color=INK)


def add_rich_text(paragraph, text, base_size=10.6, color=INK):
    # Minimal Markdown emphasis parser for bold spans and inline code.
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Menlo", size=max(base_size - 1, 8), color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size, color=color)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_GRAY)
    cell.text = ""
    first = True
    for line in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line if line else " ")
        set_font(run, name="Courier New", size=8.0, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(section):
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("GENERAL-AI-HUMAN-HOMEWORK-001  |  Page ")
    set_font(run, size=8, color=MUTED)
    add_page_field(p)


def add_title_block(doc):
    # Named override: customer_pack title furniture over compact_reference_guide.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("KAIRALI AI METHOD  |  BEGINNER HOMEWORK PACK")
    set_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Two Required Workers. One Optional Worker.")
    set_font(run, size=31, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Personal Work Memory + Daily Email EA + Drive Master Index, with optional Saturday LinkedIn")
    set_font(run, size=13.5, color=GRAY)

    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    values = [
        ("WHO", "Meeting attendee without a named homework page"),
        ("REQUIRED", "Personal Memory + Email EA + Full Drive Index"),
        ("HUMAN ROLE", "Mission, login, approval and final judgment"),
        ("PROOF", "Visible local report + evidence + validator pass"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], values):
        shade_cell(cell, PALE_GOLD)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(label + "\n")
        set_font(label_run, size=8.5, color=GOLD, bold=True)
        value_run = p.add_run(value)
        set_font(value_run, size=10.2, color=INK, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)


def logical_markdown_lines(raw_lines):
    """Join Markdown-wrapped prose and list continuations into logical blocks."""
    output = []
    current = None
    in_code = False

    def flush():
        nonlocal current
        if current is not None:
            output.append(current)
            current = None

    for raw in raw_lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush()
            output.append(line)
            in_code = not in_code
            continue
        if in_code:
            output.append(line)
            continue
        if not line.strip():
            flush()
            output.append("")
            continue

        stripped = line.strip()
        starts_block = bool(
            re.match(r"^(#{1,3})\s+", stripped)
            or re.match(r"^-\s+", stripped)
            or re.match(r"^\d+\.\s+", stripped)
            or stripped.startswith("> ")
        )

        if starts_block:
            if current and current.startswith("> ") and stripped.startswith("> "):
                current += " " + stripped[2:]
            else:
                flush()
                current = stripped
            continue

        if current is None:
            current = stripped
        elif re.match(r"^(#{1,3})\s+", current):
            flush()
            current = stripped
        else:
            current += " " + stripped

    flush()
    return output


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.right_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    add_header_footer(section)
    configure_styles(doc)
    add_title_block(doc)

    lines = logical_markdown_lines(SOURCE.read_text(encoding="utf-8").splitlines())
    # Title and subtitle are rendered by the custom customer_pack opening block.
    index = 3 if len(lines) > 2 and lines[0].startswith("# ") and lines[2].startswith("## ") else 0
    if index:
        lines = lines[3:]

    in_code = False
    code_lines = []
    list_kind = None
    list_num_id = None

    for position, raw in enumerate(lines):
        line = raw.rstrip()
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            list_kind = None
            list_num_id = None
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            list_kind = None
            list_num_id = None
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph(style=f"Heading {level}")
            add_rich_text(p, heading.group(2), base_size={1:16, 2:13, 3:12}[level], color=BLUE if level < 3 else DARK_BLUE)
            list_kind = None
            list_num_id = None
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        number = re.match(r"^(\d+)\.\s+(.+)$", line)
        if bullet or number:
            kind = "bullet" if bullet else "number"
            if list_kind != kind:
                start_value = int(number.group(1)) if number else 1
                list_num_id = add_numbering_definition(doc, kind, start_value=start_value)
                list_kind = kind
            p = doc.add_paragraph()
            apply_numbering(p, list_num_id)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.20
            following = next((candidate for candidate in lines[position + 1:] if candidate), "")
            if number and following.startswith("**DONE WHEN:**"):
                p.paragraph_format.keep_with_next = True
            add_rich_text(p, bullet.group(1) if bullet else number.group(2))
            continue

        list_kind = None
        list_num_id = None
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.20
        if line.startswith("> "):
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(9)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PALE_BLUE)
            p_pr.append(shd)
            add_rich_text(p, line[2:], color=NAVY)
        elif line.startswith("**DONE WHEN:**"):
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(9)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PALE_GOLD)
            p_pr.append(shd)
            add_rich_text(p, line, color=NAVY)
        else:
            add_rich_text(p, line)

    doc.core_properties.title = "Two Required Workers. One Optional Worker."
    doc.core_properties.subject = "Beginner Personal Work Memory, Daily Email EA, Drive Master Index and optional Saturday LinkedIn homework"
    doc.core_properties.author = "Kairali AI Method"
    doc.core_properties.keywords = "Codex, AI human, personal work memory, email EA, JSONL, drive index, LinkedIn message assistant, beginner"
    mark_table_headers(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
